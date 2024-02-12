from django.shortcuts import render
from django.http import HttpResponse
from django.shortcuts import render, redirect
from plagiarismchecker.algorithm import main
from django.forms import inlineformset_factory
from django import forms
from django.contrib.auth.forms import UserCreationForm
from django import forms
from .forms import CreateUserForm;
from docx import *
from plagiarismchecker.algorithm import fileSimilarity
import PyPDF2 
from .models import *
from django.contrib import messages
from django.contrib.auth import authenticate,login as authLogin,logout
from django.contrib.auth.decorators import login_required



    
def register(request):
    if request.user.is_authenticated:
        return redirect('home')
    else:
        form = CreateUserForm()

        if request.method == 'POST':
            form = CreateUserForm(request.POST) 
            if form.is_valid():      
                    form.save()    
                    user = form.cleaned_data.get('username')
                    messages.success(request,'Account was created for' + user)
                    return redirect('login')

        context={'form':form}
        return render(request,'pc/register.html',context)           
                     
def login(request):
     if request.user.is_authenticated:
         return redirect('home')
     else:
        if request.method == 'POST':
            username=request.POST.get('username')
            password= request.POST.get('password')
            user = authenticate(request, username = username, password = password)

            if user is not None:
                authLogin(request,user)
                return redirect('plagiarism-check-mainpage')
            else:
                messages.info(request, 'Username OR password is incorrect')
            
            
        context={}
        return render(request,'pc/login.html',context)



             

        
       

# Create your views here.
#home
def logoutUser(request):
    logout(request)
    return redirect('login')


@login_required(login_url='login')
def home(request):
    return render(request, 'pc/index.html') 



#web search(Text)
def test(request):
    print("request is welcome test")
    print(request.POST['q'])  
    
    if request.POST['q']: 
        percent,link = main.findSimilarity(request.POST['q'])
        percent = round(percent,2)
    print("Output.....................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

#web search file(.txt, .docx)
def filetest(request):
    value = ''    
    print(request.FILES['docfile'])
    if str(request.FILES['docfile']).endswith(".txt"):
        value = str(request.FILES['docfile'].read())

    elif str(request.FILES['docfile']).endswith(".docx"):
        document = Document(request.FILES['docfile'])
        for para in document.paragraphs:
            value += para.text

    elif str(request.FILES['docfile']).endswith(".pdf"):
        # creating a pdf file object 
        pdfFileObj = open(request.FILES['docfile'], 'rb') 

        # creating a pdf reader object 
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj) 

        # printing number of pages in pdf file 
        print(pdfReader.numPages) 

        # creating a page object 
        pageObj = pdfReader.getPage(0) 

        # extracting text from page 
        print(pageObj.extractText()) 

        # closing the pdf file object 
        pdfFileObj.close() 


    percent,link = main.findSimilarity(value)
    percent= round(percent, 2) 
    print("Output...................!!!!!!!!",percent,link)
    return render(request, 'pc/index.html',{'link': link, 'percent': percent})

#text compare

@login_required(login_url='login')
def fileCompare(request):
    return render(request, 'pc/doc_compare.html') 

#two text compare(Text)
def twofiletest1(request):
    print("Submiited text for 1st and 2nd")
    print(request.POST['q1'])
    print(request.POST['q2'])

    if request.POST['q1'] != '' and request.POST['q2'] != '': 
        print("Got both the texts")
        result = fileSimilarity.findFileSimilarity(request.POST['q1'],request.POST['q2'])
    result = round(result,2)    
    print("Output>>>>>>>>>>>>>>>>>>>>!!!!!!!!",result)
    return render(request, 'pc/doc_compare.html',{'result': result})
    

#two text compare(.txt, .docx)
def twofilecompare1(request):
    value1 = ''
    value2 = ''
    if (str(request.FILES['docfile1'])).endswith(".txt") and (str(request.FILES['docfile2'])).endswith(".txt"):
        value1 = str(request.FILES['docfile1'].read())
        value2 = str(request.FILES['docfile2'].read())

    elif (str(request.FILES['docfile1'])).endswith(".docx") and (str(request.FILES['docfile2'])).endswith(".docx"):
        document = Document(request.FILES['docfile1'])
        for para in document.paragraphs:
            value1 += para.text
        document = Document(request.FILES['docfile2'])
        for para in document.paragraphs:
            value2 += para.text

    result2 = fileSimilarity.findFileSimilarity(value1,value2)
    result2 = round(result2, 2)  
    print("Output..................!!!!!!!!",result2)
    return render(request, 'pc/doc_compare.html',{'result2': result2})
